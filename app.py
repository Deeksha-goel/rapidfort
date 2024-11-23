from flask import Flask, request, render_template, send_file
from werkzeug.utils import secure_filename
from docx import Document
from reportlab.pdfgen import canvas
import os

# Initialize Flask app
app = Flask(__name__)

# Configure upload and converted folders
app.config['UPLOAD_FOLDER'] = 'uploads/'
app.config['CONVERTED_FOLDER'] = 'converted_pdfs/'

# Create folders if they don't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['CONVERTED_FOLDER'], exist_ok=True)

# Route: Homepage
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Check if file is in request
        if 'file' not in request.files:
            return "No file part", 400
        
        file = request.files['file']
        
        # Check if a file is selected
        if file.filename == '':
            return "No selected file", 400
        
        # Check for .docx file and process it
        if file and file.filename.endswith('.docx'):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)

            # Extract metadata and convert to PDF
            metadata = extract_metadata(filepath)
            pdf_path = convert_to_pdf(filepath)
            return render_template('result.html', metadata=metadata, pdf_path=pdf_path)
    
    return render_template('index.html')

# Function: Extract Metadata
def extract_metadata(filepath):
    doc = Document(filepath)
    metadata = {
        "Word Count": sum(len(p.text.split()) for p in doc.paragraphs),
        "Paragraphs": len(doc.paragraphs),
    }
    return metadata

# Function: Convert to PDF
def convert_to_pdf(filepath):
    doc = Document(filepath)
    output_path = os.path.join(app.config['CONVERTED_FOLDER'], "converted.pdf")
    pdf = canvas.Canvas(output_path)

    # Initial position for text
    x, y = 50, 800

    # Add text from Word document paragraphs to PDF
    for paragraph in doc.paragraphs:
        pdf.drawString(x, y, paragraph.text)
        y -= 20  # Move down for the next line

        # Check if a new page is needed
        if y < 50:  # Threshold to avoid writing at the bottom of the page
            pdf.showPage()
            y = 800  # Reset y position for new page

    # Save the PDF file
    pdf.save()
    return output_path

# Route: Download the converted PDF
@app.route('/download/<path:filename>', methods=['GET'])
def download_file(filename):
    return send_file(filename, as_attachment=True)

# Main entry point
if __name__ == '__main__':
    app.run(debug=True)



